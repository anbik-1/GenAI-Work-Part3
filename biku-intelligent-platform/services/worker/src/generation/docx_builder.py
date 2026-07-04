"""Build a Genese-branded .docx from generated content sections.
If a custom template exists in S3 for the document type, use it as the base.
Otherwise fall back to the programmatic Genese template."""
import io
import os
import boto3
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shared import PROPOSAL_SECTIONS, SOW_SECTIONS, CASE_STUDY_SECTIONS

# Genese brand colors (used in fallback template)
GENESE_BLUE = RGBColor(0x00, 0x4E, 0x96)
GENESE_ORANGE = RGBColor(0xF5, 0x7C, 0x00)
GENESE_DARK = RGBColor(0x1A, 0x1A, 0x2E)
SECTION_TITLE_MAP = {
    "executive_summary": "Executive Summary",
    "problem_statement": "Problem Statement",
    "proposed_solution": "Proposed Solution",
    "architecture": "Solution Architecture",
    "team": "Our Team",
    "timeline": "Project Timeline",
    "investment": "Investment",
    "next_steps": "Next Steps",
    "project_overview": "Project Overview",
    "scope_of_work": "Scope of Work",
    "deliverables": "Deliverables",
    "assumptions_and_exclusions": "Assumptions and Exclusions",
    "project_team": "Project Team",
    "timeline_and_milestones": "Timeline and Milestones",
    "pricing": "Pricing",
    "terms_and_conditions": "Terms and Conditions",
    "client_overview": "Client Overview",
    "challenge": "The Challenge",
    "solution": "Our Solution",
    "architecture_overview": "Architecture Overview",
    "results": "Results and Outcomes",
    "key_takeaways": "Key Takeaways",
}


def _get_template_from_s3(document_type: str) -> bytes | None:
    """Try to download a custom template from S3. Returns None if not found."""
    from ..core.config import get_settings
    settings = get_settings()
    s3_key = f"templates/{document_type}/template.docx"
    try:
        s3 = boto3.client("s3", region_name=settings.aws_region)
        response = s3.get_object(Bucket=settings.documents_bucket, Key=s3_key)
        print(f"[docx_builder] Using custom template: {s3_key}")
        return response["Body"].read()
    except Exception:
        return None


def _embed_arch_diagram(doc: Document, png_bytes: bytes) -> None:
    """Embed the architecture PNG into the document."""
    import tempfile
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
        f.write(png_bytes)
        tmp_path = f.name
    try:
        run.add_picture(tmp_path, width=Inches(6.0))
    except Exception as e:
        print(f"[docx_builder] Could not embed diagram: {e}")
    finally:
        import os
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
    cap = doc.add_paragraph("Figure: Proposed AWS Architecture")
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _fill_custom_template(template_bytes: bytes, sections_content: dict,
                           client_name: str, engagement_type: str,
                           sources: list | None,
                           arch_png_bytes: bytes | None = None) -> bytes:
    """
    Use ONLY the styles, fonts, colors and page layout from the uploaded template.
    All existing body content is cleared — Claude's generated sections replace it.
    """
    doc = Document(io.BytesIO(template_bytes))

    # ── Step 1: Clear all body paragraphs (keep styles/theme) ────────────────
    # Remove body paragraphs but preserve document styles, page setup, and headers/footers
    body = doc.element.body
    # Remove all paragraphs and tables from body (keep sectPr — page settings)
    import lxml.etree as etree
    children_to_remove = []
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):  # paragraphs, tables, structured docs
            children_to_remove.append(child)
    for child in children_to_remove:
        body.remove(child)

    # ── Step 2: Add document title block ─────────────────────────────────────
    doc_type_label = "PROPOSAL"
    # Try to detect doc type from sections
    if any(k in sections_content for k in ('scope_of_work', 'deliverables')):
        doc_type_label = "STATEMENT OF WORK"
    elif any(k in sections_content for k in ('challenge', 'results')):
        doc_type_label = "CASE STUDY"

    # Type label
    type_para = doc.add_paragraph()
    type_run = type_para.add_run(doc_type_label)
    type_run.bold = True
    type_run.font.size = Pt(11)
    type_run.font.color.rgb = GENESE_ORANGE
    type_para.paragraph_format.space_before = Pt(12)

    # Client name as heading
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(client_name)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = GENESE_BLUE

    # Engagement type
    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(engagement_type.replace("_", " ").title())
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = GENESE_DARK

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Prepared: {datetime.now().strftime('%B %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ── Step 3: Add Claude's generated sections ───────────────────────────────
    for section_key, content in sections_content.items():
        if section_key in ("parse_error", "content"):
            _add_section(doc, "Generated Content", str(content))
            continue
        title = SECTION_TITLE_MAP.get(section_key, section_key.replace("_", " ").title())
        _add_section(doc, title, str(content))
        if section_key in ("architecture", "architecture_overview") and arch_png_bytes:
            _embed_arch_diagram(doc, arch_png_bytes)

    if sources:
        _add_sources(doc, sources)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx(
    sections_content: dict,
    document_type: str,
    client_name: str,
    engagement_type: str,
    sources: list[dict] | None = None,
    arch_png_bytes: bytes | None = None,
    template_name: str | None = None,
    plain_text_instructions: str | None = None,
) -> bytes:
    """
    Build a .docx from generated sections.

    template_name resolution order:
      - 'plain_text'  → minimal unstyled document (no branding, no header/footer)
      - any other str → treated as document_type for S3 template lookup
      - None          → check S3 for a template matching document_type, else default Genese template
    """
    # Plain-text format: skip all branding
    if template_name == "plain_text":
        return _build_plain_text_docx(sections_content, client_name, engagement_type,
                                       sources, arch_png_bytes, plain_text_instructions)

    # Custom S3 template lookup
    s3_key_type = template_name if template_name else document_type
    template_bytes = _get_template_from_s3(s3_key_type)
    if template_bytes:
        return _fill_custom_template(template_bytes, sections_content,
                                      client_name, engagement_type, sources, arch_png_bytes)
    return _build_default_docx(sections_content, document_type, client_name,
                                engagement_type, sources, arch_png_bytes)


def _build_plain_text_docx(
    sections_content: dict,
    client_name: str,
    engagement_type: str,
    sources: list[dict] | None = None,
    arch_png_bytes: bytes | None = None,
    plain_text_instructions: str | None = None,
) -> bytes:
    """
    Minimal, unstyled .docx — plain headings and body text only.
    No Genese branding, no header/footer, no colours.
    """
    doc = Document()

    # Use default margins but no header/footer decorations
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Top note
    note = doc.add_paragraph("Generated by Genese Proposal AI \u2014 Plain Text Format")
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()  # spacer

    # Format instructions (if provided)
    if plain_text_instructions and plain_text_instructions.strip():
        instructions_para = doc.add_paragraph(f"Format Instructions: {plain_text_instructions.strip()}")
        instructions_para.runs[0].italic = True
        instructions_para.runs[0].font.size = Pt(10)
        instructions_para.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        doc.add_paragraph()  # spacer

    # Document title
    title = doc.add_heading(client_name, level=1)
    # Remove any default colour so it renders in plain black
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    sub = doc.add_paragraph(engagement_type.replace("_", " ").title())
    sub.runs[0].font.size = Pt(11)
    doc.add_paragraph()  # spacer

    # Content sections
    for section_key, content in sections_content.items():
        if section_key in ("parse_error", "content"):
            heading = doc.add_heading("Generated Content", level=1)
        else:
            title_text = SECTION_TITLE_MAP.get(section_key, section_key.replace("_", " ").title())
            heading = doc.add_heading(title_text, level=1)

        # Strip any colour styling from the heading runs
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        for para_text in str(content).split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.size = Pt(11)

        # Embed architecture diagram if available
        if section_key in ("architecture", "architecture_overview") and arch_png_bytes:
            _embed_arch_diagram(doc, arch_png_bytes)

        doc.add_paragraph()  # spacer between sections

    # Sources
    if sources:
        src_heading = doc.add_heading("References", level=1)
        for run in src_heading.runs:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        for source in sources:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(source.get("title", ""))
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(f"\n{source.get('url', '')}").font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_default_docx(
    sections_content: dict,
    document_type: str,
    client_name: str,
    engagement_type: str,
    sources: list[dict] | None = None,
    arch_png_bytes: bytes | None = None,
) -> bytes:
    """Programmatic Genese-branded .docx (used when no custom template uploaded)."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Header with Genese branding
    _add_header(doc, document_type, client_name)

    # Title
    _add_title(doc, document_type, client_name, engagement_type)

    # Document sections
    for section_key, content in sections_content.items():
        if section_key in ("parse_error", "content"):
            _add_section(doc, "Generated Content", str(content))
            continue
        title = SECTION_TITLE_MAP.get(section_key, section_key.replace("_", " ").title())
        _add_section(doc, title, str(content))
        # Embed architecture diagram after the architecture section
        if section_key in ("architecture", "architecture_overview") and arch_png_bytes:
            _embed_arch_diagram(doc, arch_png_bytes)

    # Sources appendix
    if sources:
        _add_sources(doc, sources)

    # Footer
    _add_footer(doc)

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_header(doc: Document, document_type: str, client_name: str):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    run = p.add_run("GENESE SOLUTION  |  Confidential")
    run.font.size = Pt(9)
    run.font.color.rgb = GENESE_BLUE
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_title(doc: Document, document_type: str, client_name: str, engagement_type: str):
    doc_type_label = {
        "proposal": "PROPOSAL",
        "sow": "STATEMENT OF WORK",
        "case_study": "CASE STUDY",
    }.get(document_type, document_type.upper())

    # Spacer
    doc.add_paragraph()

    # Document type
    p = doc.add_paragraph()
    run = p.add_run(doc_type_label)
    run.font.size = Pt(11)
    run.font.color.rgb = GENESE_ORANGE
    run.bold = True
    run.font.all_caps = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Client name as main title
    p = doc.add_heading(client_name, level=1)
    p.runs[0].font.color.rgb = GENESE_BLUE

    # Engagement type subtitle
    p = doc.add_paragraph()
    run = p.add_run(engagement_type.replace("_", " ").title())
    run.font.size = Pt(12)
    run.font.color.rgb = GENESE_DARK
    run.italic = True

    # Date
    p = doc.add_paragraph()
    run = p.add_run(f"Prepared: {datetime.now().strftime('%B %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer


def _add_section(doc: Document, title: str, content: str):
    # Section heading
    h = doc.add_heading(title, level=2)
    h.runs[0].font.color.rgb = GENESE_BLUE

    # Content paragraphs
    for para_text in content.split("\n"):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(11)

    doc.add_paragraph()  # spacer between sections


def _add_sources(doc: Document, sources: list[dict]):
    doc.add_heading("References", level=2)
    for source in sources:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(source.get("title", ""))
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(f"\n{source.get('url', '')}").font.size = Pt(9)


def _add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    run = p.add_run("© Genese Solution. This document is confidential and intended solely for the named recipient.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
