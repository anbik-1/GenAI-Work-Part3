"""Build a Genese-branded .docx from generated content sections."""
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shared import PROPOSAL_SECTIONS, SOW_SECTIONS, CASE_STUDY_SECTIONS

# Genese brand colors
GENESE_BLUE = RGBColor(0x00, 0x4E, 0x96)    # #004E96
GENESE_ORANGE = RGBColor(0xF5, 0x7C, 0x00)  # #F57C00
GENESE_DARK = RGBColor(0x1A, 0x1A, 0x2E)    # #1A1A2E
SECTION_TITLE_MAP = {
    "executive_summary": "Executive Summary",
    "problem_statement": "Problem Statement",
    "proposed_solution": "Proposed Solution",
    "architecture": "Solution Architecture",
    "team": "Our Team",
    "timeline": "Project Timeline",
    "investment": "Investment",
    "next_steps": "Next Steps",
    "project_overview": "Project Overview",
    "scope_of_work": "Scope of Work",
    "deliverables": "Deliverables",
    "assumptions_and_exclusions": "Assumptions and Exclusions",
    "project_team": "Project Team",
    "timeline_and_milestones": "Timeline and Milestones",
    "pricing": "Pricing",
    "terms_and_conditions": "Terms and Conditions",
    "client_overview": "Client Overview",
    "challenge": "The Challenge",
    "solution": "Our Solution",
    "architecture_overview": "Architecture Overview",
    "results": "Results and Outcomes",
    "key_takeaways": "Key Takeaways",
}


def build_docx(
    sections_content: dict,
    document_type: str,
    client_name: str,
    engagement_type: str,
    sources: list[dict] | None = None,
) -> bytes:
    """
    Build a Genese-branded .docx document from generated sections.
    Returns bytes that can be uploaded to S3.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Header with Genese branding
    _add_header(doc, document_type, client_name)

    # Title
    _add_title(doc, document_type, client_name, engagement_type)

    # Document sections
    for section_key, content in sections_content.items():
        if section_key in ("parse_error", "content"):
            # Fallback — dump raw content
            _add_section(doc, "Generated Content", str(content))
            continue
        title = SECTION_TITLE_MAP.get(section_key, section_key.replace("_", " ").title())
        _add_section(doc, title, str(content))

    # Sources appendix
    if sources:
        _add_sources(doc, sources)

    # Footer
    _add_footer(doc)

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_header(doc: Document, document_type: str, client_name: str):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    run = p.add_run("GENESE SOLUTION  |  Confidential")
    run.font.size = Pt(9)
    run.font.color.rgb = GENESE_BLUE
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_title(doc: Document, document_type: str, client_name: str, engagement_type: str):
    doc_type_label = {
        "proposal": "PROPOSAL",
        "sow": "STATEMENT OF WORK",
        "case_study": "CASE STUDY",
    }.get(document_type, document_type.upper())

    # Spacer
    doc.add_paragraph()

    # Document type
    p = doc.add_paragraph()
    run = p.add_run(doc_type_label)
    run.font.size = Pt(11)
    run.font.color.rgb = GENESE_ORANGE
    run.bold = True
    run.font.all_caps = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Client name as main title
    p = doc.add_heading(client_name, level=1)
    p.runs[0].font.color.rgb = GENESE_BLUE

    # Engagement type subtitle
    p = doc.add_paragraph()
    run = p.add_run(engagement_type.replace("_", " ").title())
    run.font.size = Pt(12)
    run.font.color.rgb = GENESE_DARK
    run.italic = True

    # Date
    p = doc.add_paragraph()
    run = p.add_run(f"Prepared: {datetime.now().strftime('%B %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer


def _add_section(doc: Document, title: str, content: str):
    # Section heading
    h = doc.add_heading(title, level=2)
    h.runs[0].font.color.rgb = GENESE_BLUE

    # Content paragraphs
    for para_text in content.split("\n"):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(11)

    doc.add_paragraph()  # spacer between sections


def _add_sources(doc: Document, sources: list[dict]):
    doc.add_heading("References", level=2)
    for source in sources:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(source.get("title", ""))
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(f"\n{source.get('url', '')}").font.size = Pt(9)


def _add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    run = p.add_run("© Genese Solution. This document is confidential and intended solely for the named recipient.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
